"""
Streamlit Agentic Consent Form Generator

Gemini drives an agentic loop, calling Dentally API tools to gather complete
patient data, then generates a custom-filled consent form .docx.
"""

import datetime
import io
import json
import os
import re
import shutil
import tempfile

import requests
import streamlit as st
from bs4 import BeautifulSoup
from docx import Document
from google import genai
from google.genai import types

# ---------------------------------------------------------------------------
# Config
# ---------------------------------------------------------------------------

BASE_URL = "https://api.dentally.co/v1"
SITE_ID  = "70596be0-e19d-43a1-86e8-e209a3c1aa92"
TEMPLATE = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                         "FAME Single Implant Template.docx")

def _secret(key):
    try:
        return st.secrets[key]
    except Exception:
        return os.environ.get(key, "")

@st.cache_resource
def _gemini():
    return genai.Client(api_key=_secret("GEMINI_API_KEY"))

def _dheaders():
    return {
        "Authorization": f"Bearer {_secret('DENTALLY_API_TOKEN')}",
        "Accept": "application/json",
        "User-Agent": "DentallyConsentForm/1.0",
    }

def _dget(path, params=None):
    r = requests.get(f"{BASE_URL}{path}", headers=_dheaders(),
                     params=params, timeout=15)
    r.raise_for_status()
    return r.json()

# ---------------------------------------------------------------------------
# Tool implementations (called by the agent)
# ---------------------------------------------------------------------------

def tool_search_patients(name: str) -> str:
    patients = _dget("/patients", {"query": name}).get("patients", [])
    if not patients:
        return "No patients found."
    return json.dumps([{
        "id": p["id"],
        "name": f"{p.get('first_name','')} {p.get('last_name','')}".strip(),
        "dob": p.get("date_of_birth"),
        "email": p.get("email_address"),
        "mobile": p.get("mobile_phone"),
    } for p in patients])


def tool_get_patient(patient_id: int) -> str:
    p = _dget(f"/patients/{patient_id}").get("patient", {})
    return json.dumps({
        "id": p.get("id"),
        "name": f"{p.get('first_name','')} {p.get('last_name','')}".strip(),
        "title": p.get("title"),
        "dob": p.get("date_of_birth"),
        "email": p.get("email_address"),
        "mobile": p.get("mobile_phone"),
        "address": ", ".join(filter(None, [
            p.get("address_line_1"), p.get("address_line_2"),
            p.get("town"), p.get("postcode"),
        ])),
        "occupation": p.get("occupation"),
        "medical_alert": p.get("medical_alert"),
        "medical_alert_text": p.get("medical_alert_text"),
        "dentist_recall_date": p.get("dentist_recall_date"),
        "hygienist_recall_date": p.get("hygienist_recall_date"),
    })


def tool_get_appointments(patient_id: int) -> str:
    appts = _dget("/appointments", {
        "patient_id": patient_id, "site_id": SITE_ID, "per_page": 100,
    }).get("appointments", [])
    if not appts:
        return "No appointments found for this patient."
    return json.dumps([{
        "id": a.get("id"),
        "start_time": a.get("start_time"),
        "duration": a.get("duration"),
        "reason": a.get("reason"),
        "notes": a.get("notes"),
        "state": a.get("state"),
    } for a in appts])


def tool_get_treatment_plans(patient_id: int) -> str:
    plans = _dget("/treatment_plans", {
        "patient_id": patient_id, "site_id": SITE_ID,
    }).get("treatment_plans", [])
    if not plans:
        return "No treatment plans found."
    return json.dumps([{
        "id": p.get("id"),
        "start_date": p.get("start_date"),
        "completed": p.get("completed"),
        "nickname": p.get("nickname"),
        "practitioner_id": p.get("practitioner_id"),
    } for p in plans])


def tool_get_treatment_plan_items(treatment_plan_id: int) -> str:
    items = _dget("/treatment_plan_items", {
        "treatment_plan_id": treatment_plan_id,
    }).get("treatment_plan_items", [])
    if not items:
        return "No items in this treatment plan."
    results = []
    for item in items:
        raw = item.get("notes") or ""
        notes = BeautifulSoup(raw, "html.parser").get_text("\n").strip() if raw else ""
        results.append({
            "id": item.get("id"),
            "nomenclature": item.get("nomenclature"),
            "teeth": item.get("teeth"),
            "notes": notes,
            "completed": item.get("completed"),
            "price": item.get("price"),
        })
    return json.dumps(results)


def tool_get_recalls(patient_id: int) -> str:
    recalls = _dget("/recalls", {"patient_id": patient_id}).get("recalls", [])
    return json.dumps(recalls) if recalls else "No recalls on record."


def tool_get_sms(patient_id: int) -> str:
    sms = _dget("/sms", {"patient_id": patient_id}).get("sms", [])
    if not sms:
        return "No SMS messages on record."
    return json.dumps([{
        "id": m.get("id"),
        "direction": m.get("direction"),
        "body": m.get("body"),
        "created_at": m.get("created_at"),
        "status": m.get("status"),
    } for m in sms])


def tool_get_patient_documents(patient_id: int) -> str:
    docs = _dget("/patient_documents", {"patient_id": patient_id}).get("patient_documents", [])
    if not docs:
        return "No documents uploaded for this patient."
    return json.dumps([{
        "id": d.get("id"),
        "name": d.get("name"),
        "description": d.get("description"),
        "file_type": d.get("file_type"),
        "created_at": d.get("created_at"),
        "url": d.get("url"),
    } for d in docs])


def tool_get_notes(patient_id: int) -> str:
    notes = _dget("/notes", {"patient_id": patient_id}).get("notes", [])
    if not notes:
        return "No notes on record."
    return json.dumps([{
        "id": n.get("id"),
        "note": n.get("note"),
        "created_at": n.get("created_at"),
        "author": n.get("author"),
    } for n in notes])


def tool_get_medical_histories(patient_id: int) -> str:
    records = _dget("/medical_histories", {"patient_id": patient_id}).get("medical_histories", [])
    if not records:
        return "No medical history forms completed."
    return json.dumps([{
        "id": r.get("id"),
        "created_at": r.get("created_at"),
        "answers": r.get("answers"),
        "conditions": r.get("conditions"),
        "medications": r.get("medications"),
        "allergies": r.get("allergies"),
    } for r in records])


def tool_get_invoices(patient_id: int) -> str:
    invoices = _dget("/invoices", {"patient_id": patient_id}).get("invoices", [])
    if not invoices:
        return "No invoices on record."
    return json.dumps([{
        "id": i.get("id"),
        "state": i.get("state"),
        "total": i.get("total"),
        "balance": i.get("balance"),
        "created_at": i.get("created_at"),
        "description": i.get("description"),
    } for i in invoices])


def tool_generate_consent_form(patient_name: str, clinical_context: str) -> str:
    """Fills the FAME template using Gemini and stores the .docx in session state."""
    today  = datetime.date.today().strftime("%d %B %Y")
    prompt = f"""You are a treatment coordinator at FAME Dentistry Ltd, Edinburgh.

Using the comprehensive patient data below, write patient-friendly text for four
sections of a dental implant consent form.

PATIENT: {patient_name}  DATE: {today}  DENTIST: Dr Ferhan Ahmed

PATIENT DATA:
{clinical_context}

Return ONLY valid JSON (no markdown, no code fences):
{{
  "clinical_explanation": "2-3 warm sentences in plain English — what brought the patient in and what was found. No abbreviations.",
  "clinical_assessment": "1-2 sentences on the clinical findings in plain English.",
  "opt_note": "One sentence on what the scans/X-rays showed.",
  "appointment_1": "2-3 sentences describing the first surgical appointment based on the treatment plan."
}}"""

    resp = _gemini().models.generate_content(model="gemini-2.5-flash", contents=prompt)
    raw  = re.sub(r'^```(?:json)?\s*|\s*```$', '', resp.text.strip(), flags=re.MULTILINE)
    content = json.loads(raw)

    address    = st.session_state.get("patient_address", "")
    docx_bytes = _build_docx(patient_name, address, content)

    st.session_state.pending_form = {
        "patient_name": patient_name,
        "content": content,
        "docx_bytes": docx_bytes,
    }
    return f"Consent form generated for {patient_name}. It is ready for download."


# ---------------------------------------------------------------------------
# Tool registry & Gemini declarations
# ---------------------------------------------------------------------------

TOOL_FN_MAP = {
    "search_patients":          tool_search_patients,
    "get_patient":              tool_get_patient,
    "get_appointments":         tool_get_appointments,
    "get_treatment_plans":      tool_get_treatment_plans,
    "get_treatment_plan_items": tool_get_treatment_plan_items,
    "get_recalls":              tool_get_recalls,
    "get_sms":                  tool_get_sms,
    "get_patient_documents":    tool_get_patient_documents,
    "get_notes":                tool_get_notes,
    "get_medical_histories":    tool_get_medical_histories,
    "get_invoices":             tool_get_invoices,
    "generate_consent_form":    tool_generate_consent_form,
}

TOOL_LABELS = {
    "search_patients":          "Searching patients",
    "get_patient":              "Fetching patient details",
    "get_appointments":         "Fetching appointments",
    "get_treatment_plans":      "Fetching treatment plans",
    "get_treatment_plan_items": "Fetching clinical notes",
    "get_recalls":              "Fetching recall schedule",
    "get_sms":                  "Fetching SMS correspondence",
    "get_patient_documents":    "Fetching uploaded documents",
    "get_notes":                "Fetching patient notes",
    "get_medical_histories":    "Fetching medical history",
    "get_invoices":             "Fetching invoices",
    "generate_consent_form":    "Generating consent form",
}

def _S(t, desc=None, **props):
    return types.Schema(type=t, description=desc,
                        properties=props or None,
                        required=list(props) if props else None)

GEMINI_TOOLS = types.Tool(function_declarations=[
    types.FunctionDeclaration(
        name="search_patients",
        description="Search Dentally for patients by name. Returns IDs needed for all other calls.",
        parameters=_S(types.Type.OBJECT, name=_S(types.Type.STRING, "Patient full or partial name")),
    ),
    types.FunctionDeclaration(
        name="get_patient",
        description="Get full demographics and medical details for a patient ID.",
        parameters=_S(types.Type.OBJECT, patient_id=_S(types.Type.INTEGER, "Dentally patient ID")),
    ),
    types.FunctionDeclaration(
        name="get_appointments",
        description="Get all appointments for a patient including reasons and clinical notes.",
        parameters=_S(types.Type.OBJECT, patient_id=_S(types.Type.INTEGER, "Dentally patient ID")),
    ),
    types.FunctionDeclaration(
        name="get_treatment_plans",
        description="Get all treatment plans for a patient. Returns plan IDs required for get_treatment_plan_items.",
        parameters=_S(types.Type.OBJECT, patient_id=_S(types.Type.INTEGER, "Dentally patient ID")),
    ),
    types.FunctionDeclaration(
        name="get_treatment_plan_items",
        description="Get all items and full clinical notes within a treatment plan. "
                    "This is where consultation notes, diagnoses, and planned procedures live.",
        parameters=_S(types.Type.OBJECT, treatment_plan_id=_S(types.Type.INTEGER, "Treatment plan ID")),
    ),
    types.FunctionDeclaration(
        name="get_recalls",
        description="Get the recall schedule for a patient.",
        parameters=_S(types.Type.OBJECT, patient_id=_S(types.Type.INTEGER, "Dentally patient ID")),
    ),
    types.FunctionDeclaration(
        name="get_sms",
        description="Get all SMS messages sent to or received from a patient. Contains outbound correspondence and patient replies.",
        parameters=_S(types.Type.OBJECT, patient_id=_S(types.Type.INTEGER, "Dentally patient ID")),
    ),
    types.FunctionDeclaration(
        name="get_patient_documents",
        description="Get all documents uploaded against a patient record (letters, referrals, scan reports, consent forms, etc.).",
        parameters=_S(types.Type.OBJECT, patient_id=_S(types.Type.INTEGER, "Dentally patient ID")),
    ),
    types.FunctionDeclaration(
        name="get_notes",
        description="Get free-text clinical notes added to the patient record.",
        parameters=_S(types.Type.OBJECT, patient_id=_S(types.Type.INTEGER, "Dentally patient ID")),
    ),
    types.FunctionDeclaration(
        name="get_medical_histories",
        description="Get completed medical history forms including conditions, medications, and allergies.",
        parameters=_S(types.Type.OBJECT, patient_id=_S(types.Type.INTEGER, "Dentally patient ID")),
    ),
    types.FunctionDeclaration(
        name="get_invoices",
        description="Get invoices for a patient including treatment costs and balances.",
        parameters=_S(types.Type.OBJECT, patient_id=_S(types.Type.INTEGER, "Dentally patient ID")),
    ),
    types.FunctionDeclaration(
        name="generate_consent_form",
        description="Generate a completed dental implant consent form .docx for the patient. "
                    "Call this only after gathering all relevant patient data. "
                    "Provide a rich clinical_context string covering all findings.",
        parameters=_S(
            types.Type.OBJECT,
            patient_name=_S(types.Type.STRING, "Patient full name"),
            clinical_context=_S(types.Type.STRING,
                "Comprehensive summary of all gathered data: demographics, clinical notes, "
                "diagnoses, X-ray findings, treatment planned, teeth involved, next steps"),
        ),
    ),
])

SYSTEM_PROMPT = """You are a dental treatment coordinator assistant at FAME Dentistry Ltd, Edinburgh.
You have direct access to the Dentally patient management system through the provided tools.

Behaviour:
- When a patient name is mentioned, immediately search for them and then proactively gather
  ALL available data: patient details, appointments, treatment plans, treatment plan items
  (clinical notes live here), SMS correspondence, uploaded documents, free-text notes,
  medical history forms, and invoices.
- Answer questions using the real data you fetched, not assumptions.
- If a source returns empty, note it briefly and move on — do not stop.
- When asked to generate a consent form, first make sure you have fetched everything, then
  call generate_consent_form with a rich clinical_context string incorporating all findings
  including any correspondence context and medical history.
- After generating a form, tell the user it is ready to download and offer to make any changes.
- You can continue chatting and calling tools as needed throughout the conversation."""

# ---------------------------------------------------------------------------
# DOCX builder
# ---------------------------------------------------------------------------

def _replace_in_para(para, old, new, exact=False):
    text = para.text
    if exact:
        if text.strip() != old.strip():
            return False
    elif old not in text:
        return False
    for run in para.runs:
        if old in run.text:
            run.text = run.text.replace(old, new, 1)
            return True
    if para.runs:
        para.runs[0].text = text.replace(old, new, 1)
        for run in para.runs[1:]:
            run.text = ""
    return True


def _apply(doc, sub_map, exact_map):
    def process(para):
        for old, new in exact_map.items():
            if _replace_in_para(para, old, new, exact=True):
                return
        for old, new in sub_map.items():
            _replace_in_para(para, old, new)

    for para in doc.paragraphs:
        process(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    process(para)
    for section in doc.sections:
        for part in (section.header, section.footer,
                     section.first_page_header, section.first_page_footer):
            try:
                for para in part.paragraphs:
                    process(para)
            except Exception:
                pass


def _build_docx(patient_name, address, content):
    today = datetime.date.today().strftime("%d %B %Y")

    exact_map = {
        "Date":            today,
        "Patient Name":    patient_name,
        "Patient Address": address or "",
        "Patient friendly explanation of clinical notes.":          content["clinical_explanation"],
        "Clinical assessment explained in patient friendly words.":  content["clinical_assessment"],
        "OPT with area marked for patient reference.":              content["opt_note"],
        "Appointment explained.":                                   content["appointment_1"],
    }
    sub_map = dict(sorted({
        "Dental Implant Treatment Proposal for:  ":
            f"Dental Implant Treatment Proposal for: {patient_name}",
        "Prepared by: Dentist Name":   "Prepared by: Dr Ferhan Ahmed",
        "Dear Patient Name,":          f"Dear {patient_name},",
        "I, Patient name, hereby consent":
            f"I, {patient_name}, hereby consent",
        "explained the Treatment Plan to patient name and given her/him":
            f"explained the Treatment Plan to {patient_name} and given her",
        "Total Cost of Appointment 1   Cost":  "Total Cost of Appointment 1   TBC",
        "Total Cost of Appointment 2   Cost":  "Total Cost of Appointment 2   TBC",
        "Total Cost of Appointment 3   Cost":  "Total Cost of Appointment 3   TBC",
        "Total Cost of Appointment 4 Cost ":   "Total Cost of Appointment 4   TBC",
        "Total Cost of Appointment 5   Cost":  "Total Cost of Appointment 5   TBC",
        "Total Cost of Appointment 6    Cost": "Total Cost of Appointment 6   TBC",
        "Total Cost of Appointment 7   Cost":  "Total Cost of Appointment 7   TBC",
    }.items(), key=lambda x: len(x[0]), reverse=True))

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        shutil.copy2(TEMPLATE, tmp.name)
        tmp_path = tmp.name
    try:
        doc = Document(tmp_path)
        _apply(doc, sub_map, exact_map)
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
    finally:
        os.unlink(tmp_path)


def _docx_to_text(docx_bytes):
    doc = Document(io.BytesIO(docx_bytes))
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())

# ---------------------------------------------------------------------------
# UI components
# ---------------------------------------------------------------------------

def render_form_card(patient_name, content, docx_bytes, key):
    today = datetime.date.today().strftime("%d %B %Y")
    st.markdown(f"""
**Patient:** {patient_name} &nbsp;|&nbsp; **Date:** {today} &nbsp;|&nbsp; **Dentist:** Dr Ferhan Ahmed

---
**Your situation**
{content["clinical_explanation"]}

**What we found**
{content["clinical_assessment"]}

**Your scans**
{content["opt_note"]}

**Appointment 1**
{content["appointment_1"]}

---
*Full document includes appointments 2–7, complications list, consent & payment sections.*
""")
    col1, col2 = st.columns(2)
    with col1:
        st.download_button(
            "⬇️ Download .docx", data=docx_bytes,
            file_name=f"{patient_name} - Implant Consent Form.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True, key=f"dl_{key}",
        )
    with col2:
        with st.expander("📋 Copy full text"):
            st.text_area("", value=_docx_to_text(docx_bytes), height=260,
                         label_visibility="collapsed", key=f"txt_{key}")

# ---------------------------------------------------------------------------
# Agentic loop
# ---------------------------------------------------------------------------

def _run_agent(user_message):
    st.session_state.gemini_history.append(
        types.Content(role="user", parts=[types.Part(text=user_message)])
    )

    response_text = ""
    calls_log     = []

    with st.chat_message("assistant"):
        status_box  = st.empty()
        text_box    = st.empty()

        for _ in range(20):  # hard cap on iterations
            try:
                resp = _gemini().models.generate_content(
                    model="gemini-2.5-flash",
                    contents=st.session_state.gemini_history,
                    config=types.GenerateContentConfig(
                        system_instruction=SYSTEM_PROMPT,
                        tools=[GEMINI_TOOLS],
                    ),
                )
            except Exception as e:
                status_box.error(f"Gemini error: {e}")
                return

            candidate = resp.candidates[0]
            parts     = candidate.content.parts
            fn_calls  = [p for p in parts if p.function_call]

            if not fn_calls:
                # Terminal text response
                response_text = "".join(p.text for p in parts if hasattr(p, "text") and p.text)
                status_box.empty()
                text_box.markdown(response_text)
                st.session_state.gemini_history.append(candidate.content)
                break

            # Show live status
            names = [p.function_call.name for p in fn_calls]
            label = " · ".join(TOOL_LABELS.get(n, n) for n in names)
            status_box.info(f"🔍 {label}…")
            calls_log.extend(names)

            # Add model turn to history
            st.session_state.gemini_history.append(candidate.content)

            # Execute each tool call and collect responses
            tool_parts = []
            for p in fn_calls:
                fc   = p.function_call
                args = {k: v for k, v in fc.args.items()}
                try:
                    result = TOOL_FN_MAP[fc.name](**args)
                except Exception as e:
                    result = f"Error calling {fc.name}: {e}"
                tool_parts.append(types.Part(
                    function_response=types.FunctionResponse(
                        name=fc.name, response={"result": result}
                    )
                ))

            # Function responses go back as a "user" turn
            st.session_state.gemini_history.append(
                types.Content(role="user", parts=tool_parts)
            )

        # Show collapsed log of API calls made
        if calls_log:
            with st.expander(f"🔧 {len(calls_log)} Dentally API calls", expanded=False):
                for c in calls_log:
                    st.caption(f"• {TOOL_LABELS.get(c, c)}")

        # Render form card if the agent generated one
        pending = st.session_state.pop("pending_form", None)
        if pending:
            st.session_state.form_version += 1
            render_form_card(
                pending["patient_name"], pending["content"],
                pending["docx_bytes"], key=st.session_state.form_version,
            )
            st.session_state.messages.append({
                "role": "assistant", "text": response_text,
                "type": "form",
                "patient_name": pending["patient_name"],
                "content":      pending["content"],
                "docx_bytes":   pending["docx_bytes"],
                "form_key":     st.session_state.form_version,
                "calls_log":    calls_log,
            })
        else:
            st.session_state.messages.append({
                "role": "assistant", "text": response_text,
                "calls_log": calls_log,
            })

# ---------------------------------------------------------------------------
# Session state init
# ---------------------------------------------------------------------------

def _init():
    defaults = dict(
        messages=[],
        gemini_history=[],   # list of types.Content — persists full conversation
        patient_address="",
        pending_form=None,
        form_version=0,
    )
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v

# ---------------------------------------------------------------------------
# App
# ---------------------------------------------------------------------------

st.set_page_config(page_title="Consent Form Generator", page_icon="🦷", layout="centered")
st.title("🦷 Consent Form Generator")

_init()

if not _secret("DENTALLY_API_TOKEN") or not _secret("GEMINI_API_KEY"):
    st.error("Add DENTALLY_API_TOKEN and GEMINI_API_KEY to Streamlit secrets.")
    st.stop()

# Re-render chat history
for msg in st.session_state.messages:
    with st.chat_message(msg["role"]):
        st.markdown(msg["text"])
        if msg.get("calls_log"):
            with st.expander(f"🔧 {len(msg['calls_log'])} Dentally API calls", expanded=False):
                for c in msg["calls_log"]:
                    st.caption(f"• {TOOL_LABELS.get(c, c)}")
        if msg.get("type") == "form":
            render_form_card(
                msg["patient_name"], msg["content"], msg["docx_bytes"],
                key=f"hist_{msg['form_key']}",
            )

# Greeting
if not st.session_state.messages:
    with st.chat_message("assistant"):
        st.markdown("Hi! Which patient would you like to generate a consent form for?")
    st.session_state.messages.append({
        "role": "assistant",
        "text": "Hi! Which patient would you like to generate a consent form for?",
    })

# Chat input
if prompt := st.chat_input("Type a message…"):
    with st.chat_message("user"):
        st.markdown(prompt)
    st.session_state.messages.append({"role": "user", "text": prompt})
    _run_agent(prompt)
